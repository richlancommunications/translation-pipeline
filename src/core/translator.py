"""
Core translation engine with multi-backend support and glossary integration.
"""

import json
import re
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class TranslationResult:
    """Result of a translation operation."""
    text: str
    confidence_score: float
    glossary_matches: List[Dict]
    word_count: int
    engine_used: str
    status: str
    metadata: Dict = None


class Translator:
    """
    Main translation engine supporting multiple backends and glossary integration.
    """
    
    def __init__(
        self,
        source_lang: str = "en",
        target_lang: str = "sw",
        glossary_path: Optional[str] = None,
        use_local_llm: bool = False,
        api_key: Optional[str] = None
    ):
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.glossary = self._load_glossary(glossary_path) if glossary_path else {}
        self.use_local_llm = use_local_llm
        self.api_key = api_key
        self.translation_cache = {}
        
        logger.info(f"Translator initialized: {source_lang} -> {target_lang}")
        if self.glossary:
            logger.info(f"Loaded glossary with {len(self.glossary)} terms")
    
    def _load_glossary(self, path: str) -> Dict:
        """Load glossary from JSON file."""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                # Convert to lookup dict for faster access
                glossary = {}
                for category, terms in data.items():
                    for term in terms:
                        source = term['source'].lower()
                        glossary[source] = {
                            'target': term['target'],
                            'context': term.get('context', ''),
                            'confidence': term.get('confidence', 1.0),
                            'alternatives': term.get('alternatives', [])
                        }
                return glossary
        except Exception as e:
            logger.error(f"Error loading glossary: {e}")
            return {}
    
    def translate_text(
        self,
        text: str,
        domain: Optional[str] = None,
        preserve_formatting: bool = True
    ) -> TranslationResult:
        """
        Translate a text string with glossary support.
        
        Args:
            text: Source text to translate
            domain: Optional domain context for better term selection
            preserve_formatting: Maintain original text formatting
            
        Returns:
            TranslationResult with translation and metadata
        """
        # Check cache first
        cache_key = f"{text[:100]}_{self.source_lang}_{self.target_lang}"
        if cache_key in self.translation_cache:
            logger.info("Using cached translation")
            return self.translation_cache[cache_key]
        
        # Step 1: Apply glossary replacements
        glossary_matches = []
        preprocessed_text, matches = self._apply_glossary(text, domain)
        glossary_matches.extend(matches)
        
        # Step 2: Translate using selected backend
        if self.use_local_llm:
            translated = self._translate_with_ollama(preprocessed_text)
            engine = "ollama"
        else:
            translated = self._translate_with_google(preprocessed_text)
            engine = "google"
        
        # Step 3: Calculate confidence score
        confidence = self._calculate_confidence(
            text, 
            translated, 
            len(glossary_matches)
        )
        
        # Step 4: Create result
        result = TranslationResult(
            text=translated,
            confidence_score=confidence,
            glossary_matches=glossary_matches,
            word_count=len(text.split()),
            engine_used=engine,
            status="success",
            metadata={
                'domain': domain,
                'glossary_coverage': len(glossary_matches) / len(text.split()) if text.split() else 0
            }
        )
        
        # Cache the result
        self.translation_cache[cache_key] = result
        
        return result
    
    def _apply_glossary(
        self, 
        text: str, 
        domain: Optional[str] = None
    ) -> Tuple[str, List[Dict]]:
        """
        Apply glossary terms to text before translation.
        Returns preprocessed text and list of matches.
        """
        matches = []
        result_text = text
        
        # Sort glossary terms by length (longest first) to handle multi-word terms
        sorted_terms = sorted(
            self.glossary.items(), 
            key=lambda x: len(x[0]), 
            reverse=True
        )
        
        for source_term, term_data in sorted_terms:
            # Check if domain matches if specified
            if domain and term_data.get('context'):
                if not term_data['context'].startswith(domain):
                    continue
            
            # Create case-insensitive pattern
            pattern = re.compile(r'\b' + re.escape(source_term) + r'\b', re.IGNORECASE)
            
            # Find matches
            for match in pattern.finditer(result_text):
                matches.append({
                    'source': match.group(),
                    'target': term_data['target'],
                    'position': match.start(),
                    'confidence': term_data['confidence'],
                    'context': term_data.get('context', '')
                })
                
                # Replace in text (for now, keep source to maintain context)
                # In production, you might want to mark these for post-processing
        
        return result_text, matches
    
    def _translate_with_ollama(self, text: str) -> str:
        """Translate using local Ollama LLM."""
        try:
            import requests
            
            prompt = f"""Translate the following text from {self.source_lang} to {self.target_lang}.
Maintain the original meaning and technical accuracy.

Text: {text}

Translation:"""
            
            response = requests.post(
                "http://localhost:11434/api/generate",
                json={
                    "model": "mistral",
                    "prompt": prompt,
                    "stream": False
                },
                timeout=120
            )
            
            if response.status_code == 200:
                result = response.json()
                translation = result.get('response', '').strip()
                logger.info("Ollama translation successful")
                return translation
            else:
                logger.error(f"Ollama error: {response.status_code}")
                return self._translate_with_google(text)  # Fallback
                
        except Exception as e:
            logger.error(f"Ollama translation failed: {e}")
            return self._translate_with_google(text)  # Fallback
    
    def _translate_with_google(self, text: str) -> str:
        """Translate using Google Translate API."""
        try:
            # Using googletrans library (free tier)
            from googletrans import Translator as GoogleTranslator
            
            gt = GoogleTranslator()
            result = gt.translate(text, src=self.source_lang, dest=self.target_lang)
            logger.info("Google Translate successful")
            return result.text
            
        except Exception as e:
            logger.error(f"Google Translate failed: {e}")
            return f"[Translation Error: {str(e)}]"
    
    def _calculate_confidence(
        self, 
        source: str, 
        translation: str, 
        glossary_matches: int
    ) -> float:
        """
        Calculate confidence score for translation.
        Factors: glossary coverage, length similarity, translation success.
        """
        if "[Translation Error" in translation:
            return 0.0
        
        # Factor 1: Glossary coverage (0-40 points)
        source_words = len(source.split())
        glossary_score = min(40, (glossary_matches / max(source_words, 1)) * 100)
        
        # Factor 2: Length similarity (0-30 points)
        length_ratio = len(translation) / max(len(source), 1)
        if 0.5 <= length_ratio <= 2.0:
            length_score = 30
        elif 0.3 <= length_ratio <= 3.0:
            length_score = 15
        else:
            length_score = 5
        
        # Factor 3: Basic validation (0-30 points)
        validation_score = 30 if translation and len(translation) > 10 else 10
        
        total_score = glossary_score + length_score + validation_score
        return round(min(100, total_score) / 100, 2)
    
    def translate_document(
        self,
        input_file: str,
        output_file: Optional[str] = None,
        domain: Optional[str] = None
    ) -> TranslationResult:
        """
        Translate a document file (DOCX, PDF, TXT).
        
        Args:
            input_file: Path to source document
            output_file: Path for translated document (auto-generated if None)
            domain: Optional domain context
            
        Returns:
            TranslationResult with document translation metadata
        """
        input_path = Path(input_file)
        
        if not input_path.exists():
            logger.error(f"File not found: {input_file}")
            return TranslationResult(
                text="",
                confidence_score=0.0,
                glossary_matches=[],
                word_count=0,
                engine_used="none",
                status="error",
                metadata={'error': 'File not found'}
            )
        
        # Extract text based on file type
        if input_path.suffix == '.docx':
            text = self._extract_from_docx(input_file)
        elif input_path.suffix == '.pdf':
            text = self._extract_from_pdf(input_file)
        elif input_path.suffix == '.txt':
            with open(input_file, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            logger.error(f"Unsupported file type: {input_path.suffix}")
            return TranslationResult(
                text="",
                confidence_score=0.0,
                glossary_matches=[],
                word_count=0,
                engine_used="none",
                status="error",
                metadata={'error': 'Unsupported file type'}
            )
        
        # Translate the text
        result = self.translate_text(text, domain=domain)
        
        # Save to output file if specified
        if output_file:
            self._save_translation(result.text, output_file, input_path.suffix)
            logger.info(f"Translation saved to {output_file}")
        
        return result
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            doc = Document(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            return ""
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ''
                for page in reader.pages:
                    text += page.extract_text()
                return text
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            return ""
    
    def _save_translation(self, text: str, output_file: str, format: str):
        """Save translation to file in appropriate format."""
        output_path = Path(output_file)
        
        if format == '.docx':
            try:
                from docx import Document
                doc = Document()
                for paragraph in text.split('\n'):
                    doc.add_paragraph(paragraph)
                doc.save(output_path)
            except Exception as e:
                logger.error(f"Error saving DOCX: {e}")
        else:
            # Default to text file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)


# Example usage
if __name__ == "__main__":
    # Initialize translator with medical glossary
    translator = Translator(
        source_lang="en",
        target_lang="sw",
        glossary_path="glossaries/medical_terms.json",
        use_local_llm=True
    )
    
    # Translate sample text
    sample_text = "The patient presents with hypertension and elevated blood sugar levels."
    result = translator.translate_text(sample_text, domain="medical")
    
    print(f"Original: {sample_text}")
    print(f"Translation: {result.text}")
    print(f"Confidence: {result.confidence_score * 100}%")
    print(f"Glossary matches: {len(result.glossary_matches)}")
    print(f"Engine: {result.engine_used}")